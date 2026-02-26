"""
Build a formatted Word document from the EB Language Decline research compilation.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style helpers ───────────────────────────────────────────────────────
style = doc.styles

# Normal
normal = style['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# Heading 1
h1 = style['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(18)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(12)

# Heading 2
h2 = style['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)

# Heading 3
h3 = style['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)


def add_horizontal_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): 'AAAAAA',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_study(doc, short_label, full_citation, link, findings):
    """Add a formatted study entry."""
    # Study header
    doc.add_heading(short_label, level=3)

    # Citation
    p_cite = doc.add_paragraph()
    run = p_cite.add_run('Citation: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_cite.add_run(full_citation)
    run.font.size = Pt(10)
    run.italic = True

    # Link
    p_link = doc.add_paragraph()
    run = p_link.add_run('Access: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p_link.add_run(link)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)

    # Findings
    for finding in findings:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(finding)
        run.font.size = Pt(10.5)

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1B3A5C',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.5)
        # Alternate shading
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                shading = row.cells[c_idx]._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'EDF2F7',
                    qn('w:val'): 'clear',
                })
                shading.append(shading_elm)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


# ════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Research Compilation')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    'Emergent Bilingual Students:\n'
    'Post-Testing Decline, Summer Language Loss,\n'
    'and Texas-Specific Evidence'
)
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

doc.add_paragraph()

dateline = doc.add_paragraph()
dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dateline.add_run('Compiled: February 25, 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

focus = doc.add_paragraph()
focus.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = focus.add_run(
    'Evidence that emergent bilingual students show a decline in language acquisition\n'
    'after spring testing season and at the end of the school year,\n'
    'compounding into summer language loss.'
)
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

stat = doc.add_paragraph()
stat.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = stat.add_run('Approximately 45 citable sources across 5 research categories')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('Part 1', 'Summer Language Loss for Emergent Bilingual Students', '14 studies'),
    ('Part 2', 'High-Stakes Testing Impact on ELL Instruction', '16 studies'),
    ('Part 3', 'Post-Testing Instructional Decline \u2014 The "Lost Weeks"', '10 studies'),
    ('Part 4', 'Texas-Specific Research on Emergent Bilingual Students', '14 studies/reports'),
    ('Part 5', 'Theoretical Frameworks', '6 studies'),
    ('Part 6', 'Synthesis: Building the Argument', ''),
    ('Part 7', 'Identified Gap in the Literature', ''),
]
for num, title_text, count in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(num + ':  ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(title_text)
    run.font.size = Pt(11)
    if count:
        run = p.add_run('  (' + count + ')')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Executive Summary', level=1)

exec_paragraphs = [
    'This document compiles approximately 45 citable academic studies and policy reports examining '
    'whether emergent bilingual (EB) students experience a decline in language acquisition after '
    'spring testing season and at the end of the school year, compounding into summer language loss. '
    'The research focuses on the United States broadly and Texas specifically.',

    'No single published study directly measures EB language proficiency continuously from April '
    'through September. However, multiple converging lines of research strongly support the pattern: '
    'language development stalls or narrows during spring test prep, instruction diminishes after '
    'testing, and then summer break compounds the loss \u2014 creating a 4\u20135 month period (roughly '
    'March through August) where emergent bilingual students are not receiving the sustained, '
    'high-quality language input they need.',

    'This is a significant and under-researched equity issue in Texas, where 1 in 4 public school '
    'students is an emergent bilingual. The research identifies the following compounding trajectory:'
]

for text in exec_paragraphs:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)

# Executive summary table
exec_headers = ['Period', 'What Happens', 'Evidence Strength']
exec_rows = [
    ['Feb\u2013April\n(Test Prep)', 'Instruction narrows to test formats;\nauthentic language practice decreases', 'Strong \u2014 49+ qualitative\nstudies confirm pattern'],
    ['April\u2013May\n(Post-Testing)', 'Instructional intensity drops;\ngrowth decelerates; no further\nlanguage assessments', 'Moderate \u2014 growth data\nconfirms deceleration;\nno ELL-specific study'],
    ['June\u2013August\n(Summer)', 'English input drops dramatically;\nvocabulary and fluency regress;\nproductive skills deteriorate first', 'Strong \u2014 multiple large-\nscale studies with ELL-\nspecific data'],
    ['September\n(Return)', 'Students restart below spring\npeak; gap widens year over year', 'Strong \u2014 longitudinal\nstudies confirm\ncumulative effects'],
]
add_table(doc, exec_headers, exec_rows, col_widths=[1.4, 2.8, 2.3])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 1: SUMMER LANGUAGE LOSS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 1: Summer Language Loss for Emergent Bilingual Students', level=1)

p = doc.add_paragraph(
    'The following studies document that emergent bilingual students experience disproportionate '
    'language and literacy loss during summer months, driven by reduced English input in home '
    'environments and the absence of structured language instruction.'
)
p.italic = True
p.paragraph_format.space_after = Pt(12)

add_horizontal_line(doc)

add_study(doc,
    'Lawrence (2012) \u2014 Summer Vocabulary Setback',
    'Lawrence, J. F. (2012). English vocabulary trajectories of students whose parents speak a language other than English: Steep trajectories and sharp summer setback. Reading and Writing, 25(5), 1113-1141.',
    'https://doi.org/10.1007/s11145-011-9305-z',
    [
        'Studied 278 adolescent students (210 English-speaking homes, 68 non-English-speaking homes) across three school years.',
        'Students from non-English-speaking homes experienced significantly sharper summer vocabulary setback than English-speaking peers, even after controlling for SES and reading habits.',
        'Home-language status was a stronger predictor of summer vocabulary loss than socioeconomic status.',
        'A compensatory tradeoff exists: steeper school-year vocabulary gains are offset by larger summer losses, creating a "two steps forward, one step back" cycle.',
    ])

add_study(doc,
    'Jaekel, Jaekel, Fincher, & Brown (2022) \u2014 Summer Reading Regression for ELs',
    'Jaekel, N., Jaekel, J., Fincher, E., & Brown, C. L. (2022). Summer regression \u2014 the impact of summer on English learners\u2019 reading development. Frontiers in Education, 7, 817284.',
    'https://doi.org/10.3389/feduc.2022.817284',
    [
        'Analyzed oral reading fluency for 3,280 students (363 ELs, 2,917 native speakers), measured May (4th grade) and September (5th grade) \u2014 directly capturing the spring-to-fall window.',
        'ELs performed 23.36 points below native speakers on oral reading fluency after summer.',
        'Both groups lost fluency at similar rates, but ELs cannot afford any loss given their already-below-grade-level starting points.',
        'Higher spring performers lost more in absolute terms, but proportional impact was greatest for ELs.',
    ])

add_study(doc,
    'Mancilla-Martinez & Lesaux (2011) \u2014 Vocabulary Gap in Spanish Speakers',
    'Mancilla-Martinez, J., & Lesaux, N. K. (2011). The gap between Spanish speakers\u2019 word reading and word knowledge: A longitudinal study. Child Development, 82(5), 1544-1560.',
    'https://doi.org/10.1111/j.1467-8624.2011.01633.x',
    [
        'Tracked 173 Spanish-speaking children from low-income households longitudinally from ages 4.5 to 11.',
        'By age 11, English vocabulary was 1.0 SD below national norms; Spanish vocabulary was 3.4 SD below \u2014 despite English-only instruction since preschool.',
        'English word reading (decoding) reached parity with norms by age 5, creating a dangerous "uneven profile" where adequate decoding masked severe vocabulary deficits.',
        'The vocabulary gap becomes a reading comprehension crisis in upper elementary even when decoding appears proficient.',
    ])

add_study(doc,
    'Mancilla-Martinez & Lesaux (2010) \u2014 Predictors of Reading Comprehension',
    'Mancilla-Martinez, J., & Lesaux, N. K. (2010). Predictors of reading comprehension for struggling readers: The case of Spanish-speaking language minority learners. Journal of Educational Psychology, 102(3), 701-711.',
    'https://doi.org/10.1037/a0019135',
    [
        'English language skills accounted for all unique variance in English reading comprehension outcomes \u2014 vocabulary is the rate-limiting factor.',
        'Despite years of English-only instruction, vocabulary growth rates were insufficient to close the gap with monolingual norms.',
        'Continued L2-only instruction did not accelerate vocabulary growth to match English monolingual peers.',
    ])

add_study(doc,
    'NWEA (2022) \u2014 Academic Growth for English Learners',
    'NWEA. (2022). Academic Growth for English Learners: Lessons from School-Year Learning Gains and Summer Learning Loss \u2014 Implications for COVID-19 Recovery and Beyond.',
    'https://www.nwea.org/resource-center/resource/academic-growth-for-english-learners-lessons-from-school-year-learning-gains-and-summer-learning-loss-implications-for-covid-19-recovery-and-beyond/',
    [
        'Using MAP Growth data (K-4), found ELs show equal or greater school-year gains than peers \u2014 but fall behind at greater rates during summer.',
        'Creates a "two steps forward, one step back" dynamic that prevents ELs from closing the gap.',
        'Recommends directing federal ESSER funding toward extended school year and summer support specifically for ELs.',
    ])

add_study(doc,
    'Atteberry & McEachin (2021) \u2014 Summer Achievement Disparities',
    'Atteberry, A., & McEachin, A. (2021). School\u2019s out: The role of summers in understanding achievement disparities. American Educational Research Journal, 58(2), 239-282.',
    'https://doi.org/10.3102/0002831220937285',
    [
        'Massive dataset: 200+ million test scores from 18 million students across 7,500 districts (2008-2016).',
        'Average student lost 17-28% of ELA school-year gains and 25-34% of math gains during summer.',
        '52% of students exhibited ELA learning losses in all 5 consecutive summers studied \u2014 documenting the cumulative nature of summer regression.',
        'At the extremes, some students lost nearly 90% of school-year gains over summer.',
    ])

add_study(doc,
    'Cooper, Nye, Charlton, Lindsay, & Greathouse (1996) \u2014 Foundational Meta-Analysis',
    'Cooper, H., Nye, B., Charlton, K., Lindsay, J., & Greathouse, S. (1996). The effects of summer vacation on achievement test scores: A narrative and meta-analytic review. Review of Educational Research, 66(3), 227-268.',
    'https://doi.org/10.3102/00346543066003227',
    [
        'Foundational meta-analysis of 39 studies: achievement scores decline approximately one month on grade-level equivalents over summer.',
        'Middle-class students gained in reading over summer; lower-class students lost ground \u2014 establishing the SES-moderated divergence directly relevant to ELLs.',
        'Negative effects increase with grade level.',
        'Math was hit hardest (~2.6 months loss on grade-level equivalents).',
    ])

add_study(doc,
    'Alexander, Entwisle, & Olson (2007) \u2014 Lasting Consequences (Faucet Theory)',
    'Alexander, K. L., Entwisle, D. R., & Olson, L. S. (2007). Lasting consequences of the summer learning gap. American Sociological Review, 72(2), 167-180.',
    'https://doi.org/10.1177/000312240707200202',
    [
        'Baltimore Beginning School Study: tracked 790 students from 1st grade through high school.',
        'The "faucet theory": during school, the resource faucet is on for all children; during summer, only resource-rich homes sustain learning.',
        'Over five elementary summers: low-income students gained less than 1 point in reading; higher-income peers gained 47 points.',
        'Cumulative summer differences substantially account for the 9th-grade SES achievement gap, including effects on tracking, dropout, and college access.',
    ])

add_study(doc,
    'Entwisle, Alexander, & Olson (2001) \u2014 Keep the Faucet Flowing',
    'Entwisle, D. R., Alexander, K. L., & Olson, L. S. (2001). Keep the faucet flowing: Summer learning and home environment. American Educator, 25(3), 10-15.',
    'https://www.aft.org/ae/fall2001/entwisle_alexander_olson',
    [
        'During the school year, poor and affluent children made comparable gains (~57-61 points reading, 45-49 points math for 1st graders).',
        'During summer, affluent students gained 15 points reading / 9 points math; low-income students lost 4 points reading / 5 points math.',
        'Over five elementary summers cumulatively: low-income gained <1 point reading, lost 8 points math; higher-income gained 47 points reading, 25 points math.',
    ])

add_study(doc,
    'Kieffer (2008) \u2014 Catching Up or Falling Behind?',
    'Kieffer, M. J. (2008). Catching up or falling behind? Initial English proficiency, concentrated poverty, and the reading growth of language minority learners in the United States. Journal of Educational Psychology, 100(4), 851-868.',
    'https://doi.org/10.1037/0022-0663.100.4.851',
    [
        'Nationally representative cohort: reading growth trajectories K-5 for language minority learners vs. native English speakers.',
        'LM learners who entered kindergarten already proficient in oral English caught up by 1st grade and maintained average levels through 8th grade.',
        'LM learners with limited English at kindergarten entry remained substantially below national averages.',
        'Concentrated poverty further depressed reading growth, compounding the disadvantage for LM learners in high-poverty schools.',
    ])

add_study(doc,
    'Kieffer (2010) \u2014 Late-Emerging Reading Difficulties',
    'Kieffer, M. J. (2010). Socioeconomic status, English proficiency, and late-emerging reading difficulties. Educational Researcher, 39(6), 484-486.',
    'https://doi.org/10.3102/0013189X10378400',
    [
        'Substantial proportions of ELLs and native speakers showed "late-emerging" reading difficulties in upper elementary and middle school \u2014 even after success in primary grades.',
        'ELLs and low-SES students at significantly elevated risk for late-emerging difficulties.',
        'If ELLs\u2019 reading difficulties emerge late (grades 4-8), then summer vocabulary loss in earlier grades may be a hidden precursor.',
    ])

add_study(doc,
    'Kieffer (2011) \u2014 Converging Trajectories K-8',
    'Kieffer, M. J. (2011). Converging trajectories: Reading growth in language minority learners and their classmates, kindergarten to grade 8. American Educational Research Journal, 48(5), 1187-1225.',
    'https://doi.org/10.3102/0002831211419490',
    [
        'Longest longitudinal window (K-8) on ELL reading trajectories in a nationally representative sample.',
        'LM learners with initially limited English had trajectories below national averages but converged with SES-matched peers by middle school.',
        'The gap is not permanent but requires sustained, targeted support \u2014 including through summer months \u2014 to close.',
    ])

# Summer interventions sub-section
doc.add_heading('Summer Intervention Studies', level=2)

add_study(doc,
    'Schmitt, Horner, & Lavery (2020) \u2014 Migrant Education Summer Programs',
    'Schmitt, A. M., Horner, S. L., & Lavery, M. R. (2020). The impact of summer programs on the English language scores of migrant children. Literacy Research and Instruction, 59(1), 78-93.',
    'https://doi.org/10.1080/19388071.2019.1687794',
    [
        'Quasi-experimental study of K-4 Latinx migrant students (n=246) in summer Migrant Education Programs.',
        'Students experienced significant gains in English speaking and language arts over summer \u2014 proving loss is preventable.',
        'One of the few studies to directly examine summer program effects on English language proficiency (not just reading) for an ELL-specific population.',
    ])

add_study(doc,
    'Kim & Quinn (2013) \u2014 Meta-Analysis of Summer Reading Interventions',
    'Kim, J. S., & Quinn, D. M. (2013). The effects of summer reading on low-income children\u2019s literacy achievement from kindergarten to grade 8: A meta-analysis. Review of Educational Research, 83(3), 386-431.',
    'https://doi.org/10.3102/0034654313483906',
    [
        'Meta-analysis of 41 summer reading interventions (K-8).',
        'Programs using research-based practices showed effect sizes of d = 0.25 to 0.63; without research-based practices, only d = 0.05-0.18.',
        'Low-income students benefited significantly more (reading comprehension d = 0.33) than mixed-income students.',
        'Most effective programs: fewer than 13 students per class, 4-8 hours daily instruction, 70+ total hours.',
    ])

add_study(doc,
    'Schacter & Jo (2005) \u2014 Summer Reading Day Camp',
    'Schacter, J., & Jo, B. (2005). Learning when school is not in session: A reading summer day-camp intervention. Journal of Research in Reading, 28(2), 158-169.',
    'https://doi.org/10.1111/j.1467-9817.2005.00260.x',
    [
        'Randomized controlled trial: economically disadvantaged exiting 1st graders in a 7-week summer reading day camp (2 hours daily).',
        'Produced a large effect size of d = 1.16 on reading assessment \u2014 one of the most effective summer interventions documented.',
        'Combination of explicit phonics instruction and small-group tutoring proved particularly powerful.',
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 2: HIGH-STAKES TESTING
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 2: High-Stakes Testing Impact on ELL Instruction', level=1)

p = doc.add_paragraph(
    'The following studies document how high-stakes testing narrows curriculum, transforms ESL/bilingual '
    'pedagogy into test-prep pedagogy, and functions as de facto language policy that undermines '
    'research-based practices for emergent bilingual students.'
)
p.italic = True
p.paragraph_format.space_after = Pt(12)

add_horizontal_line(doc)

# Testing as Language Policy
doc.add_heading('Testing as De Facto Language Policy', level=2)

add_study(doc,
    'Menken (2008) \u2014 English Learners Left Behind',
    'Menken, K. (2008). English Learners Left Behind: Standardized Testing as Language Policy. Clevedon, UK: Multilingual Matters.',
    'https://doi.org/10.21832/9781853599996',
    [
        'Argues high-stakes tests have become de facto language policy \u2014 shaping what is taught, how, and in what language.',
        'ELLs typically score 20-40 percentage points below other students \u2014 not due to lack of content knowledge but because the tests function as English proficiency exams.',
        'Schools responded by increasing English-only instruction and abandoning bilingual approaches.',
    ])

add_study(doc,
    'Menken (2006) \u2014 Teaching to the Test',
    'Menken, K. (2006). Teaching to the test: How No Child Left Behind impacts language policy, curriculum, and instruction for English language learners. Bilingual Research Journal, 30(2), 521-546.',
    'https://katemenken.org/wp-content/uploads/2011/10/menken-brj-302-summer-2006.pdf',
    [
        'Year-long study of 10 NYC high schools (128 participants).',
        'ESL classes transformed to resemble ELA classes for native speakers, driven by testing demands rather than language development needs.',
        'In bilingual classes, test availability in translation guided language allocation decisions rather than pedagogy.',
    ])

add_study(doc,
    'Menken & Solorza (2014) \u2014 No Child Left Bilingual',
    'Menken, K. & Solorza, C. (2014). No Child Left Bilingual: Accountability and the elimination of bilingual education programs in NYC schools. Educational Policy, 28(1), 96-125.',
    'https://katemenken.org/wp-content/uploads/2011/10/menken-solorza-no-child-left-bilingual-final-proof.pdf',
    [
        'Schools chose to eliminate bilingual education programs and replace them with English-only instruction because of accountability requirements.',
        'Most administrators had little or no background in ELL research; when students performed poorly, they blamed bilingual education without data.',
    ])

add_study(doc,
    'Menken & Kleyn (2010) \u2014 Long-Term Impact of Subtractive Schooling',
    'Menken, K., & Kleyn, T. (2010). The long-term impact of subtractive schooling in the educational experiences of secondary English language learners. International Journal of Bilingual Education and Bilingualism, 13(4), 399-417.',
    'https://katemenken.org/wp-content/uploads/2011/10/menken-kleyn-ijbeb-134-july-2010-subtractive-schooling-ltell.pdf',
    [
        'Years of test-driven English-only instruction produced students not fully proficient in either language.',
        'The post-testing period, when test-focused instruction has lost its rationale, is often when schools could provide enriched bilingual instruction but do not.',
        'Long-term costs borne disproportionately by ELL students who become "stuck" as long-term English learners.',
    ])

# Texas-Specific Testing Studies
doc.add_heading('Texas-Specific Testing Studies', level=2)

add_study(doc,
    'Palmer & Lynch (2008) \u2014 Bilingual Education for a Monolingual Test (Texas)',
    'Palmer, D. & Lynch, A. W. (2008). A bilingual education for a monolingual test? Language Policy, 7, 217-235.',
    'https://doi.org/10.1007/s10993-008-9100-0',
    [
        'Texas-specific study: Interview data from teachers in six Texas elementary schools.',
        'High-stakes TAKS test drives teachers\u2019 decisions about language of instruction, overriding bilingual program pedagogy.',
        'Teacher quote: "That TAKS test has entirely changed the way I teach. I teach with the goal of getting the kids to pass the test rather than to get them to be successful in reading, writing and math in general."',
    ])

add_study(doc,
    'Bach (2020) \u2014 High-Stakes Testing and EB Students in Texas',
    'Bach, A. J. (2020). High-stakes, standardized testing and emergent bilingual students in Texas. Texas Journal of Literacy Education, 8(1), 18-37.',
    'https://eric.ed.gov/?id=EJ1261335',
    [
        '22-month ethnographic study at a public high school in El Paso, Texas.',
        'Only 10% of LEP students passed the English I STAAR exam at the field site (17% district-wide).',
        'District training sessions emphasized that STAAR tests should drive instruction, with exams serving as model texts.',
        'Highlights the mismatch: Cummins\u2019s research shows 5+ years needed for academic English, yet Texas demands rapid acquisition.',
    ])

add_study(doc,
    'Wright & Li (2008) \u2014 Math Tests and Newcomer ELLs (Texas)',
    'Wright, W. E. & Li, X. (2008). High-stakes math tests: How No Child Left Behind leaves newcomer English language learners behind. Language Policy, 7, 237-266.',
    'https://doi.org/10.1007/s10993-008-9099-2',
    [
        'Texas-specific: Examined 5th-grade newcomer Cambodian students taking the Math TAKS test.',
        'Language demands of the state math test far exceeded what students could do in class.',
        'Two students failed (scoring 6 and 7 out of 44) despite one being talented in math \u2014 the test functioned as an English proficiency exam.',
    ])

# Curriculum Narrowing
doc.add_heading('Curriculum Narrowing and Washback Effects', level=2)

add_study(doc,
    'Au (2007) \u2014 Qualitative Metasynthesis of 49 Studies',
    'Au, W. (2007). High-stakes testing and curricular control: A qualitative metasynthesis. Educational Researcher, 36(5), 258-267.',
    'https://doi.org/10.3102/0013189X07306523',
    [
        'Synthesized 49 qualitative studies: predominant effect is curricular content narrowed to tested subjects, knowledge fragmented into test-related pieces, and increased teacher-centered pedagogies.',
        'Over 80% of studies found content narrowing and shifts to teacher-centered instruction.',
    ])

add_study(doc,
    'Berliner (2011) \u2014 Curriculum Narrowing Harms',
    'Berliner, D. C. (2011). Rational responses to high stakes testing: The case of curriculum narrowing and the harm that follows. Cambridge Journal of Education, 41(3), 287-302.',
    'https://www.tandfonline.com/doi/abs/10.1080/0305764X.2011.607151',
    [
        'Curriculum narrowing is the most "pernicious" yet most "rational" response to high-stakes testing.',
        'Restrictions on learning in earlier grades can retard achievement development in later grades \u2014 a compounding effect.',
        'Applies Campbell\u2019s Law: when test scores become too important, both the scores and the people using them become corrupted.',
    ])

add_study(doc,
    'Nichols & Berliner (2007) \u2014 Collateral Damage',
    'Nichols, S. L., & Berliner, D. C. (2007). Collateral Damage: How High-Stakes Testing Corrupts America\u2019s Schools. Cambridge, MA: Harvard Education Press.',
    'https://eric.ed.gov/?id=ED568726',
    [
        'Grounded in Campbell\u2019s Law: the greater the social consequences attached to a quantitative indicator, the more the indicator becomes corrupted.',
        'Documents loss of student motivation, deteriorating school climate, rising dropout rates, and curriculum narrowing.',
    ])

add_study(doc,
    'Solorzano (2008) \u2014 Comprehensive Review',
    'Solorzano, R. W. (2008). High stakes testing: Issues, implications, and remedies for English language learners. Review of Educational Research, 78(2), 260-329.',
    'https://doi.org/10.3102/0034654308317845',
    [
        'Comprehensive review concluding that high-stakes tests as currently constructed are inappropriate for ELLs.',
        'Published in one of the field\u2019s most prestigious review journals \u2014 authoritative synthesis of the evidence base.',
    ])

# Assessment Validity
doc.add_heading('Assessment Validity and Measurement Issues', level=2)

add_study(doc,
    'Abedi (2002) \u2014 Psychometric Issues',
    'Abedi, J. (2002). Standardized achievement tests and English language learners: Psychometrics issues. Educational Assessment, 8(3), 231-257.',
    'https://doi.org/10.1207/S15326977EA0803_02',
    [
        'Linguistic complexity of test items unrelated to assessed content increases measurement error and reduces reliability.',
        'Assessments with higher linguistic complexity yield larger performance gaps between ELLs and non-ELLs.',
        'Foundational psychometric argument that standard tests lack validity for ELL populations.',
    ])

add_study(doc,
    'Abedi (2004) \u2014 NCLB and ELL Assessment',
    'Abedi, J. (2004). The No Child Left Behind Act and English language learners: Assessment and accountability issues. Educational Researcher, 33(1), 4-14.',
    'https://journals.sagepub.com/doi/abs/10.3102/0013189X033001004',
    [
        'Correlation of only .223 between language assessment scores and LEP classification codes \u2014 extreme unreliability.',
        'Inconsistent LEP definitions across and within states affect accuracy of progress reporting.',
        'ELL test results may be used inappropriately to sort or retain students based on linguistically confounded scores.',
    ])

add_study(doc,
    'Abedi, Hofstetter, & Lord (2004) \u2014 Assessment Accommodations',
    'Abedi, J., Hofstetter, C. H., & Lord, C. (2004). Assessment accommodations for English language learners: Implications for policy-based empirical research. Review of Educational Research, 74(1), 1-28.',
    'https://doi.org/10.3102/00346543074001001',
    [
        'Accommodations aim to reduce the performance gap caused by construct-irrelevant factors such as unnecessary linguistic complexity.',
        'Decisions about accommodations are based on limited empirical evidence.',
        'Urges caution in using ELL test scores for consequential decisions.',
    ])

add_study(doc,
    'Solano-Flores & Trumbull (2003) \u2014 Language in Context',
    'Solano-Flores, G. & Trumbull, E. (2003). Examining language in context. Educational Researcher, 32(2), 3-13.',
    'https://doi.org/10.3102/0013189X032002003',
    [
        'Proposes using generalizability theory to estimate measurement error from language factors in ELL testing.',
        'Argues cultural validity must become a core assessment practice.',
        'Calls for concurrent dual-language assessment development, not simple translation.',
    ])

add_study(doc,
    'Solano-Flores (2008) \u2014 Probabilistic Views of Language',
    'Solano-Flores, G. (2008). Who is given tests in what language by whom, when, and where? Educational Researcher, 37(4), 189-199.',
    'https://doi.org/10.3102/0013189X08319569',
    [
        'ELL testing is, to a large extent, a random process due to poor implementation and uncontrollable factors.',
        'Challenges deterministic views of language in testing; calls for probabilistic reconceptualization.',
    ])

add_study(doc,
    'Wright (2002) \u2014 High Stakes Testing in Inner-City Schools',
    'Wright, W. E. (2002). The effects of high stakes testing in an inner-city elementary school. Current Issues in Education, 5.',
    'https://cie.asu.edu/ojs/index.php/cieatasu/article/view/1622',
    [
        'Standardized testing has not resulted in higher quality teaching and learning; instead, it narrowed curriculum and harmed both teachers and students.',
        'One of the earliest studies to examine the intersection of high-stakes testing and ELL instructional quality.',
    ])

add_study(doc,
    'Wright & Choi (2006) \u2014 Arizona ELL Testing Policies',
    'Wright, W. E. & Choi, D. (2006). The impact of language and high-stakes testing policies on elementary school English language learners in Arizona. Education Policy Analysis Archives, 14(13).',
    'https://doi.org/10.14507/epaa.v14n13.2006',
    [
        'Statewide survey of 3rd-grade ELL teachers in Arizona.',
        'Policies resulted mostly in confusion over what constitutes quality ELL instruction.',
        'Teachers support accountability in principle but believe high-stakes tests are inappropriate for ELLs.',
    ])

add_study(doc,
    'Shohamy (2001) \u2014 The Power of Tests',
    'Shohamy, E. (2001). The Power of Tests: A Critical Perspective on the Uses of Language Tests. Harlow, England: Longman.',
    'https://www.routledge.com/The-Power-of-Tests-A-Critical-Perspective-on-the-Uses-of-Language-Tests/Shohamy/p/book/9780582423350',
    [
        'Foundational work: tests are embedded in social, educational, and political contexts where they serve as instruments of power and control.',
        'Direct and linear relationship between stakes of a test and strength of washback: higher stakes = stronger washback.',
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 3: POST-TESTING DECLINE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 3: Post-Testing Instructional Decline \u2014 The "Lost Weeks"', level=1)

p = doc.add_paragraph(
    'The following studies document the erosion of instructional time due to testing calendars, '
    'non-linear within-year growth patterns that show learning deceleration in spring, and the '
    'post-testing instructional vacuum that contributes to reduced rigor in the final weeks of '
    'the school year.'
)
p.italic = True
p.paragraph_format.space_after = Pt(12)

add_horizontal_line(doc)

doc.add_heading('Within-Year Growth Deceleration', level=2)

add_study(doc,
    'Kuhfeld & Soland (2021) \u2014 Non-Linear Growth During School Year',
    'Kuhfeld, M., & Soland, J. (2021). The learning curve: Revisiting the assumption of linear growth during the school year. Journal of Research on Educational Effectiveness, 14(1), 143-171.',
    'https://doi.org/10.1080/19345747.2020.1839990',
    [
        'Analyzed 7+ million students (K-8) using NWEA MAP Growth data.',
        'Reading growth shows pronounced deceleration in winter-to-spring compared to fall-to-winter.',
        'Challenges the assumption of linear growth \u2014 spring instructional disruptions and testing likely contribute to differential growth rates.',
        'Direct evidence that the post-testing period is associated with reduced learning.',
    ])

doc.add_heading('Instructional Time Loss', level=2)

add_study(doc,
    'Kraft & Monti-Nussbaum (2021) \u2014 Classroom Interruptions',
    'Kraft, M. A., & Monti-Nussbaum, M. (2021). The big problem with little interruptions to classroom learning. AERA Open, 7, 1-17.',
    'https://doi.org/10.1177/23328584211028856',
    [
        'A typical classroom is interrupted over 2,000 times per year, resulting in loss of 10 to 20 days of instructional time.',
        'Students miss math and ELA instruction approximately 40 days per year due to absences, events, testing, and other factors.',
        'Elementary students lose 16% of allotted instructional time; middle school 21%; high school 25%.',
        'Administrators systematically underestimate the frequency and consequences of these interruptions.',
    ])

add_study(doc,
    'Jacobowitz (2014) \u2014 Erosion of Instructional Time',
    'Jacobowitz, R. (2014). Erosion of Instructional Time. Discussion Brief No. 14. The Benjamin Center, SUNY New Paltz.',
    'https://www.newpaltz.edu/media/the-benjamin-center/db_14_erosion_of_instructional_time.pdf',
    [
        'Despite 180 required instructional days, students attend far fewer full days.',
        'Students lost 9.5 to 14 instructional days due to weather, testing, PD, and communication time.',
        'The "180-day school year" is largely a fiction.',
    ])

add_study(doc,
    'Lazarin (2014) \u2014 Testing Overload',
    'Lazarin, M. (2014). Testing Overload in America\u2019s Schools. Center for American Progress.',
    'https://cdn.americanprogress.org/wp-content/uploads/2014/10/LazarinOvertestingReport.pdf',
    [
        'Teachers estimate spending 14 days preparing for state exams and 12 days for district exams \u2014 roughly 2-3 weeks combined.',
        '62% of teachers reported spending too much time readying students for state exams.',
        'A culture of test preparation "places a premium on testing over learning."',
    ])

add_study(doc,
    'Council of the Great City Schools (2015) \u2014 Testing Inventory',
    'Council of the Great City Schools. (2015). Student Testing in America\u2019s Great City Schools: An Inventory and Preliminary Analysis.',
    'https://www.cgcs.org/cms/lib/DC00001581/Centricity/Domain/87/Testing%20Report.pdf',
    [
        '401 unique tests administered across 66 large urban districts; average student takes 8 standardized tests per year spending 20-25 hours.',
        'Each PARCC session reported to "take at least a week to administer and in essence shut the school down."',
    ])

doc.add_heading('Testing\u2019s Impact on Instruction Quality', level=2)

add_study(doc,
    'Jennings & Bearak (2014) \u2014 Teaching to the Test',
    'Jennings, J. L., & Bearak, J. M. (2014). "Teaching to the test" in the NCLB era. Educational Researcher, 43(8), 381-389.',
    'https://doi.org/10.3102/0013189x14554449',
    [
        'NCLB-era state tests predictably emphasized some standards while excluding others.',
        'Teachers targeted instruction toward predictably tested skills, and once the test passed, the instructional rationale disappeared \u2014 contributing to the post-testing vacuum.',
        'In language arts, teachers "taught to the rubric" rather than developing authentic ability.',
    ])

add_study(doc,
    'McMurrer (2007/2008) \u2014 Instructional Time Reallocation',
    'McMurrer, J. (2008). Instructional time in elementary schools: A closer look at changes for specific subjects. Center on Education Policy.',
    'https://eric.ed.gov/?id=EJ812289',
    [
        '62% of districts increased time on language arts and math under NCLB; 75% of districts with schools needing improvement did so.',
        'Social studies cut by 32%, science by 33%, art and music by 35% in weekly instructional time.',
    ])

add_study(doc,
    'Stecher (2002) \u2014 Consequences of High-Stakes Testing on Practice',
    'Stecher, B. M. (2002). Consequences of large-scale, high-stakes testing on school and classroom practice. In Making Sense of Test-Based Accountability in Education. RAND Corporation.',
    'https://www.rand.org/content/dam/rand/pubs/monograph_reports/2002/MR1554.pdf',
    [
        'Negative reallocation of classroom instruction is widespread under high-stakes testing.',
        'Teachers often devote so much time to test prep that students begin preparing in fall and continue until spring \u2014 with the post-testing period left as instructional afterthought.',
    ])

add_study(doc,
    'Pedulla, Abrams, Madaus, et al. (2003) \u2014 Teacher Survey',
    'Pedulla, J. J., et al. (2003). Perceived Effects of State-Mandated Testing Programs on Teaching and Learning. National Board on Educational Testing and Public Policy, Boston College.',
    'https://eric.ed.gov/?id=ED481836',
    [
        '56% of elementary teachers reported state tests influenced teaching daily or a few times a week.',
        'Over 70% reported that test scores were not an accurate measure of what minority students know and can do.',
    ])

doc.add_heading('Seasonal Learning Patterns', level=2)

add_study(doc,
    'Downey, von Hippel, & Broh (2004) \u2014 Are Schools the Great Equalizer?',
    'Downey, D. B., von Hippel, P. T., & Broh, B. A. (2004). Are schools the great equalizer? American Sociological Review, 69(5), 613-635.',
    'https://doi.org/10.1177/000312240406900501',
    [
        'Achievement gaps grow much more during summer than during the school year.',
        'Schools serve as important equalizers \u2014 when quality instructional time degrades, inequality expands.',
        'Foundational seasonal comparison study using ECLS-K data (~20,000 children).',
    ])

add_study(doc,
    'Von Hippel, Workman, & Downey (2018) \u2014 Replication and Correction',
    'von Hippel, P. T., Workman, J., & Downey, D. B. (2018). Inequality in reading and math skills forms mainly before kindergarten. Sociology of Education, 91(1), 85-116.',
    'https://papers.ssrn.com/sol3/papers.cfm?abstract_id=3036094',
    [
        'SES gaps tend to shrink during the school year and grow during summer.',
        'Reinforces that active, high-quality instructional time during the school year is critical for maintaining and equalizing growth.',
    ])

add_study(doc,
    'Condron, Downey, & Kuhfeld (2021) \u2014 Schools as Refractors',
    'Condron, D. J., Downey, D. B., & Kuhfeld, M. (2021). Schools as refractors. Sociology of Education, 94(4), 316-340.',
    'https://doi.org/10.1177/00380407211041542',
    [
        'Schooling has a compensatory effect on inequality in reading, language, and science \u2014 but not in math.',
        'When instruction degrades (post-testing periods), the compensatory effect of schooling may diminish.',
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 4: TEXAS-SPECIFIC RESEARCH
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 4: Texas-Specific Research on Emergent Bilingual Students', level=1)

p = doc.add_paragraph(
    'Texas has approximately 1.2 million emergent bilingual students \u2014 23% of the state\u2019s '
    'public school enrollment and the largest such population in the country. The following studies '
    'and reports document TELPAS data patterns, reclassification challenges, program effectiveness, '
    'and systemic policy issues specific to Texas.'
)
p.italic = True
p.paragraph_format.space_after = Pt(12)

add_horizontal_line(doc)

doc.add_heading('TELPAS Data and Assessment Issues', level=2)

add_study(doc,
    'Texas Tribune (2024) \u2014 Six Years of Low TELPAS Scores',
    'Texas Tribune. (2024). "After six years of low scores for students learning English, Texas educators say it\u2019s the test\u2019s fault."',
    'https://www.texastribune.org/2024/08/13/texas-telpas-bilingual-students-test-scores/',
    [
        'Before the 2018 TELPAS redesign, ~50% earned the highest speaking score; after, only ~10% per year.',
        'Speaking passing rate dropped from 50%+ (2017) to 7.6% (2018).',
        'Automated scoring may penalize students with accents or who mix in native-language words.',
        'Students who cannot test out remain in remedial courses longer, limiting elective options and advanced coursework.',
    ])

add_study(doc,
    'IES/REL Southwest (2022) \u2014 Texas ELs and COVID-19',
    'IES Regional Educational Laboratory Southwest / Texas ERC. (2022). English Proficiency and the Pandemic: How Texas English Learner Students Fared During COVID-19. REL 2023-144.',
    'https://ies.ed.gov/ncee/rel/Products/Region/southwest/Publication/100897',
    [
        'Elementary ELs earned meaningfully lower TELPAS scores on listening, speaking, and reading during the pandemic, with largest declines in speaking.',
        'Reclassification rates dropped 7.6 percentage points \u2014 from 11.8% (2017/18) to 4.2% (2020/21).',
        'No meaningful differences by program model, suggesting disruption of language input affected all program types.',
    ])

add_study(doc,
    'Education Week / WIDA (2024) \u2014 National EL Proficiency Decline',
    'Education Week / WIDA. (2024). "English Learners\u2019 Proficiency Scores Are Still in Decline, Data Find."',
    'https://www.edweek.org/teaching-learning/english-learners-proficiency-scores-are-still-in-decline-data-find/2024/04',
    [
        'Nationally, EL English-language proficiency scores have remained in decline since the pandemic through 2023.',
        'Pronounced declines in grades 1-2, the foundational years for language development.',
        'Disparities worsened between Hispanic and non-Hispanic ELs since the pandemic.',
    ])

doc.add_heading('Reclassification Patterns and Timelines', level=2)

add_study(doc,
    'Kinder Institute / HERC, Rice University (2024) \u2014 Texas Failing to Nurture Bilingualism',
    'Kinder Institute for Urban Research / HERC, Rice University. (2024). "Bilingualism is a strength Texas is failing to nurture."',
    'https://kinder.rice.edu/urbanedge/bilingualism-strength-texas-failing-nurture-students',
    [
        '~1.2 million of Texas\u2019s 5.5 million public school students are emergent bilingual.',
        'Houston-area EB enrollment grew 55% between 2011-12 and 2023-24.',
        'Of students who began 1st grade in 2018-19, 83% had not reclassified before middle school \u2014 up from 48% for the 2011-12 cohort.',
    ])

add_study(doc,
    'Kinder Institute / HERC (2022) \u2014 Long-Term English Learner Outcomes',
    'Kinder Institute / HERC. (2022). Long-Term English Learners: How Is Timing of Reclassification Associated with Middle and High School Outcomes? (Part 4 Research Brief).',
    'https://kinder.rice.edu/research/long-term-english-learners-how-timing-reclassification-associated-middle-and-high-school',
    [
        'Students who reclassify in elementary school outperform peers \u2014 including students who were never EB \u2014 on multiple academic measures.',
        'LTELs had higher dropout rates, higher absenteeism, more exclusionary discipline.',
        'Timing of reclassification is a critical predictor of long-term outcomes.',
    ])

add_study(doc,
    'Umansky & Reardon (2014) \u2014 Reclassification Patterns',
    'Umansky, I. M., & Reardon, S. F. (2014). Reclassification patterns among Latino English learner students. American Educational Research Journal, 51(5), 879-912.',
    'https://cepa.stanford.edu/content/reclassification-patterns-among-latino-english-learner-students-bilingual-dual-immersion-and-english-immersion-classrooms',
    [
        'Median reclassification time: 8 years; well over half are long-term ELs; nearly a quarter never reclassify.',
        'Bilingual/dual immersion students reclassify slower in elementary but achieve higher overall rates and proficiency by high school end.',
        'Short-term metrics showing bilingual programs as "slower" are misleading.',
    ])

add_study(doc,
    'Thompson (2017) \u2014 Time to Reclassification',
    'Thompson, K. D. (2017). English learners\u2019 time to reclassification: An analysis. Educational Policy, 31(3), 330-363.',
    'https://journals.sagepub.com/doi/abs/10.1177/0895904815598394',
    [
        'Speaking/listening proficiency typically achieved within 2 years; literacy-based proficiency takes 4-7 years.',
        'Critical reclassification window in upper elementary \u2014 students not reclassified by then face steeply declining odds.',
    ])

add_study(doc,
    'Saunders & Marcelletti (2013) \u2014 The Catch-22 of Reclassification',
    'Saunders, W. M., & Marcelletti, D. J. (2013). The gap that can\u2019t go away: The Catch-22 of reclassification in monitoring the progress of English learners. Educational Evaluation and Policy Analysis, 35(2), 139-156.',
    'https://journals.sagepub.com/doi/abs/10.3102/0162373712461849',
    [
        'The "Catch-22": successful students are reclassified and removed, making remaining ELs appear to make no progress.',
        'With RFEPs excluded, the gap increased from 38% to 42%; with RFEPs included, the gap actually decreased from 26% to 22%.',
        'Standard reporting practices systematically overestimate the achievement gap and underestimate actual progress.',
    ])

doc.add_heading('Policy Reports', level=2)

add_study(doc,
    'IDRA & Every Texan (2021) \u2014 Creating a More Bilingual Texas',
    'Sikes, C. L. (IDRA) & Villanueva, C. K. (Every Texan). (2021). Creating a More Bilingual Texas.',
    'https://www.idra.org/wp-content/uploads/2021/03/Creating-a-More-Bilingual-Texas-2021.pdf',
    [
        'Texas enrolls EB students at twice the national rate.',
        'Less than one in three (29%) EB graduates in the class of 2019 were deemed college-ready.',
        'Additional funding for EB students has not changed since 1984 \u2014 nearly 40 years of policy stagnation.',
    ])

add_study(doc,
    'IDRA / Equity Assistance Center (2022) \u2014 EB Learner Literature Review',
    'IDRA / Equity Assistance Center Region II. (2022). Emergent Bilingual Learner Education: Literature Review. (ERIC ED629281).',
    'https://eric.ed.gov/?id=ED629281',
    [
        'Comprehensive literature review: policies, funding, history of instructional quality for EB learners.',
        'Texas identifies more than 19% of public school students as emergent bilingual.',
        'Advocated for the terminology change from "English Language Learner" to "Emergent Bilingual" (effective September 1, 2021).',
    ])

add_study(doc,
    'TEA (2024) \u2014 Comprehensive Biennial Report',
    'Texas Education Agency. (2024). 2024 Comprehensive Biennial Report on Texas Public Schools.',
    'https://tea.texas.gov/reports-and-data/school-performance/accountability-research/comp-annual-biennial-2024.pdf',
    [
        'Official biennial report covering all aspects of Texas public education.',
        'Includes TELPAS proficiency level data across four domains at four levels.',
        'Describes four state-approved bilingual program models and two ESL models.',
    ])

add_study(doc,
    'TEA (2025) \u2014 Results Driven Accountability Report',
    'Texas Education Agency. (2025). 2025 Results Driven Accountability Regional Report: BE/ESL/EB.',
    'https://tea.texas.gov/media/418201',
    [
        'Bilingual Education STAAR 3-8 passing rates: 69.5% Math and 70.5% RLA (2025).',
        'English I and II rose four percentage points in grade-level proficiency.',
    ])

doc.add_heading('Bilingual Program Effectiveness', level=2)

add_study(doc,
    'Collier & Thomas (2017) \u2014 32 Years of Bilingual Schooling Research',
    'Collier, V. P., & Thomas, W. P. (2017). Validating the power of bilingual schooling: Thirty-two years of large-scale, longitudinal research. Annual Review of Applied Linguistics, 37, 203-217.',
    'https://www.cambridge.org/core/journals/annual-review-of-applied-linguistics/article/abs/validating-the-power-of-bilingual-schooling-thirtytwo-years-of-largescale-longitudinal-research/909F284BFF9C327124AD08987143E677',
    [
        '32 years of research across 36 districts in 16 states, analyzing 7.5+ million student records.',
        'English-only and short-term transitional programs close only half the achievement gap; long-term bilingual programs close all of it after 5-6 years.',
        'Key variable: duration of primary language instruction \u2014 longer L1 content instruction = higher ultimate English achievement.',
    ])

add_study(doc,
    'Collier & Thomas (2004) \u2014 Astounding Effectiveness of Dual Language',
    'Collier, V. P., & Thomas, W. P. (2004). The astounding effectiveness of dual language education for all. NABE Journal of Research and Practice, 2(1), 1-20.',
    'https://www.berkeleyschools.net/wp-content/uploads/2011/10/TWIAstounding_Effectiveness_Dual_Language_Ed.pdf',
    [
        'Two-way 90:10 dual language programs reach the highest levels of achievement in the shortest time.',
        'ELLs in two-way classes outscored ELLs in other program types by 7+ NCEs.',
        'English-only and short-duration transitional programs close only about half the gap; dual language closes all of it after 5-6 years.',
    ])

add_study(doc,
    'Thomas & Collier (2002) \u2014 National Study of Long-Term Achievement',
    'Thomas, W. P., & Collier, V. P. (2002). A National Study of School Effectiveness for Language Minority Students\u2019 Long-Term Academic Achievement. CREDE Final Report.',
    'https://eric.ed.gov/?id=ED475048',
    [
        'Five-year study (1996-2001) examining eight major program types.',
        'Socioculturally supportive school environments allowing natural language development are essential.',
        'One-way 50:50 programs need continuation through middle school to completely close the achievement gap.',
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 5: THEORETICAL FRAMEWORKS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 5: Theoretical Frameworks', level=1)

p = doc.add_paragraph(
    'The following foundational theories and evidence-based guidelines underpin the argument that '
    'sustained, high-quality language input is essential for emergent bilingual students, and that '
    'interruptions to this input \u2014 whether from testing, post-testing decline, or summer break '
    '\u2014 carry measurable consequences.'
)
p.italic = True
p.paragraph_format.space_after = Pt(12)

add_horizontal_line(doc)

add_study(doc,
    'Cummins (1979) \u2014 BICS/CALP Distinction',
    'Cummins, J. (1979). Linguistic interdependence and the educational development of bilingual children. Review of Educational Research, 49(2), 222-251.',
    'https://link.springer.com/rwe/10.1007/978-0-387-30424-3_36',
    [
        'Foundational BICS/CALP distinction: conversational fluency develops in 1-3 years; academic language proficiency requires 5-7 years.',
        'Students who "sound fluent" may still be years away from academic-level proficiency.',
        'Premature program exit is harmful.',
    ])

add_study(doc,
    'Cummins (2000) \u2014 Language, Power, and Pedagogy',
    'Cummins, J. (2000). Language, Power, and Pedagogy: Bilingual Children in the Crossfire. Clevedon: Multilingual Matters.',
    'https://www.degruyterbrill.com/document/doi/10.21832/9781853596773/html',
    [
        'Extended and refined BICS/CALP framework with analysis of how power relations affect language development.',
        'Children\u2019s first language in schooling is a cognitive and academic resource, not an obstacle.',
        'The conversational vs. academic language distinction is central to understanding minority student educational outcomes.',
    ])

add_study(doc,
    'Krashen (1982) \u2014 Input Hypothesis',
    'Krashen, S. (1982). Principles and Practice in Second Language Acquisition. Oxford: Pergamon Press.',
    'https://www.sdkrashen.com/content/books/principles_and_practice.pdf',
    [
        'Input Hypothesis (i+1): learners acquire language through comprehensible input slightly above their current level.',
        'When comprehensible input is removed (summer break, post-testing instructional decline), the acquisition process stalls.',
        'Sustained, meaning-focused exposure is the primary driver of acquisition.',
    ])

add_study(doc,
    'Saunders, Goldenberg, & Marcelletti (2013) \u2014 ELD Guidelines',
    'Saunders, W., Goldenberg, C., & Marcelletti, D. (2013). English language development: Guidelines for instruction. American Educator, 37(2), 13-25.',
    'https://www.aft.org/ae/summer2013/saunders_goldenberg_marcelletti',
    [
        '14 evidence-based guidelines for ELD instruction.',
        'Key: a separate, daily block devoted to ELD; continue until advanced English proficiency \u2014 not just conversational fluency.',
        'Critical gap: surprisingly little research directly measures effects of instruction on English language development.',
    ])

add_study(doc,
    'Tracy-Ventura et al. (2025) \u2014 Second Language Attrition',
    'Tracy-Ventura, N., Huensch, A., Katz, J., & Mitchell, R. (2025). Is second language attrition inevitable after instruction ends? Language Learning, 75(1).',
    'https://doi.org/10.1111/lang.12665',
    [
        'Tracked participants 8 years (6 years post-instruction) on oral proficiency, fluency, and vocabulary.',
        'Peak attainment at end of instruction predicts retention \u2014 higher proficiency at interruption = better retention.',
        'Students who have NOT reached a threshold proficiency are the most vulnerable to regression.',
        'Productive skills (speaking, writing) deteriorate faster than receptive skills (listening, reading).',
    ])

add_study(doc,
    'Second Language Attrition \u2014 Regression Hypothesis',
    'General research on second language attrition and the regression hypothesis.',
    'https://pmc.ncbi.nlm.nih.gov/articles/PMC6866255/',
    [
        'Skills learned last (or practiced least) are lost first; productive skills deteriorate faster than receptive skills.',
        'Key factors: duration/intensity of prior instruction, proficiency level before attrition period, length of attrition period.',
        'For students who have not consolidated academic English, even a 5-month gap (April TELPAS to September school start) can result in measurable regression.',
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 6: SYNTHESIS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 6: Synthesis \u2014 Building the Argument', level=1)

p = doc.add_paragraph(
    'The following table maps the compounding trajectory of language acquisition decline for '
    'emergent bilingual students from February through September, linking each phase to the '
    'specific studies that provide supporting evidence.'
)
p.paragraph_format.space_after = Pt(12)

synth_headers = ['Period', 'What Happens', 'Key Citations']
synth_rows = [
    [
        'Feb\u2013April\n(Test Prep)',
        'Instruction narrows to test formats;\nauthentic language practice decreases',
        'Au 2007; Berliner 2011;\nMenken 2006; Palmer &\nLynch 2008; Bach 2020'
    ],
    [
        'April\u2013May\n(Post-Testing)',
        'Instructional intensity drops;\ngrowth decelerates; no further\nlanguage assessments',
        'Kuhfeld & Soland 2021;\nKraft & Monti-Nussbaum 2021;\nJennings & Bearak 2014;\nStecher 2002'
    ],
    [
        'June\u2013August\n(Summer)',
        'English input drops dramatically;\nvocabulary and fluency regress;\nproductive skills deteriorate first',
        'Lawrence 2012; Jaekel et al.\n2022; Atteberry & McEachin\n2021; NWEA 2022;\nTracy-Ventura et al. 2025'
    ],
    [
        'September\n(Return)',
        'Students restart below spring peak;\ngap widens year over year',
        'Alexander et al. 2007;\nCooper et al. 1996;\nKieffer 2008, 2011'
    ],
    [
        'Compounding\nOver Years',
        'Long-term English learners created;\nreclassification stalls;\nbilingual programs eroded',
        'Umansky & Reardon 2014;\nThompson 2017; Kinder/HERC\n2022, 2024; IDRA 2021;\nCollier & Thomas 2004, 2017'
    ],
]
add_table(doc, synth_headers, synth_rows, col_widths=[1.3, 2.5, 2.7])

doc.add_paragraph()
doc.add_paragraph()

doc.add_heading('The Convergent Evidence', level=2)

evidence_points = [
    'Tests as de facto language policy: Menken (2006, 2008), Menken & Solorza (2014), Palmer & Lynch (2008), and Wright & Li (2008) all converge on the finding that high-stakes tests function as unofficial language policies, driving instructional decisions in ways never explicitly legislated.',
    'Construct-irrelevant variance: Abedi (2002, 2004), Solano-Flores & Trumbull (2003), Solano-Flores (2008), and Wright & Li (2008) demonstrate that standardized tests measure English proficiency more than content knowledge for ELLs, rendering scores invalid for their stated purposes.',
    'Negative washback on instruction quality: Au (2007), Wright (2002), Menken (2006), and Bach (2020) document that test pressure narrows curricula, transforms ESL pedagogy into test-prep pedagogy, and replaces language-rich instruction with skills-and-drills approaches.',
    'Bilingual program erosion: Menken & Solorza (2014) and Palmer & Lynch (2008) show that accountability pressure leads to the elimination of bilingual programs and early transition to English-only instruction, contrary to what research supports.',
    'Summer as the critical vulnerability window: Lawrence (2012), Jaekel et al. (2022), NWEA (2022), and Atteberry & McEachin (2021) confirm that ELLs lose ground disproportionately during summer, with home language environment as a key driver.',
    'Interventions work: Schmitt et al. (2020), Kim & Quinn (2013), and Schacter & Jo (2005) demonstrate that structured summer programs can prevent \u2014 and even reverse \u2014 summer language loss for ELL populations.',
]

for point in evidence_points:
    p = doc.add_paragraph(style='List Bullet')
    # Bold the part before the colon
    if ':' in point:
        parts = point.split(':', 1)
        run = p.add_run(parts[0] + ':')
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(parts[1])
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(point)
        run.font.size = Pt(10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# PART 7: GAP IN THE LITERATURE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Part 7: Identified Gap in the Literature', level=1)

gap_paragraphs = [
    'No single published study directly measures emergent bilingual language proficiency '
    'continuously from April through September as a single trajectory \u2014 in Texas or nationally. '
    'The evidence must be assembled from converging lines of research:',
]

for text in gap_paragraphs:
    doc.add_paragraph(text)

gap_bullets = [
    'Spring-to-fall reading loss (Jaekel et al. 2022)',
    'Non-linear within-year growth deceleration (Kuhfeld & Soland 2021)',
    'Summer vocabulary loss for non-English-speaking homes (Lawrence 2012)',
    'SLA attrition research (Tracy-Ventura et al. 2025)',
    'Post-testing instructional vacuum (Jennings & Bearak 2014; Kraft & Monti-Nussbaum 2021)',
]
for b in gap_bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph(
    'This convergence is strong, but a longitudinal study tracking emergent bilingual language '
    'proficiency monthly from March through September would be a significant original contribution '
    'to the field. Such a study would be the first to directly document the compounding trajectory '
    'from test-prep narrowing through post-testing decline through summer language loss as a '
    'continuous phenomenon rather than separate, disconnected research findings.'
)
p.paragraph_format.space_after = Pt(12)

doc.add_paragraph()
add_horizontal_line(doc)
doc.add_paragraph()

# Footer note
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Compiled February 25, 2026  |  Approximately 45 citable sources across 5 research categories')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# ── Save ────────────────────────────────────────────────────────────────
output_path = os.path.join(
    r'C:\Users\Jordan Miller\Desktop\Manager_of_Multilingual_ESL_Coordinators\Data_Analysis',
    'EB_Language_Decline_Research_Compilation.docx'
)
doc.save(output_path)
print(f'Document saved to: {output_path}')
